"""
Handles file uploads: extracts and preprocesses text from
PDF, CSV, XLSX, DOCX, TXT, and image files (PNG/JPG).

Extraction strategy:
  - PDF / Images  → AWS Textract (high-accuracy OCR for invoices/scanned docs)
  - CSV           → Python csv (aggregated summary)
  - XLSX/XLS      → openpyxl
  - DOCX          → python-docx
  - TXT           → plain decode
"""
import io
import csv
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = (
    ".pdf", ".csv", ".xlsx", ".xls", ".docx", ".txt",
    ".png", ".jpg", ".jpeg",
)


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Route file to the appropriate extractor."""
    fn = filename.lower()
    if fn.endswith(".pdf"):
        return _extract_from_pdf_textract(file_bytes, filename)
    elif fn.endswith((".png", ".jpg", ".jpeg")):
        return _extract_from_image_textract(file_bytes, filename)
    elif fn.endswith(".csv"):
        return _extract_from_csv(file_bytes)
    elif fn.endswith((".xlsx", ".xls")):
        return _extract_from_excel(file_bytes)
    elif fn.endswith(".docx"):
        return _extract_from_docx(file_bytes)
    elif fn.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        try:
            return file_bytes.decode("utf-8", errors="replace")
        except Exception:
            return ""


# ── AWS Textract ──────────────────────────────────────────────────────────────

def _get_textract_client():
    """Build a Textract client using the same credentials as Bedrock."""
    import boto3
    import os
    from pathlib import Path
    from dotenv import load_dotenv

    env_path = Path(__file__).resolve().parent.parent / ".env"
    load_dotenv(dotenv_path=env_path, override=True)

    kwargs = {"region_name": os.getenv("AWS_DEFAULT_REGION") or os.getenv("AWS_REGION", "us-east-1")}
    key = os.getenv("AWS_ACCESS_KEY_ID", "")
    secret = os.getenv("AWS_SECRET_ACCESS_KEY", "")
    token = os.getenv("AWS_SESSION_TOKEN", "")
    if key:
        kwargs["aws_access_key_id"] = key
    if secret:
        kwargs["aws_secret_access_key"] = secret
    if token:
        kwargs["aws_session_token"] = token
    return boto3.client("textract", **kwargs)


def _textract_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Call Textract detect_document_text on raw bytes.
    Returns concatenated text from all LINE blocks.
    Falls back to local extraction on any error.
    """
    try:
        client = _get_textract_client()
        response = client.detect_document_text(
            Document={"Bytes": file_bytes}
        )
        lines = [
            block["Text"]
            for block in response.get("Blocks", [])
            if block["BlockType"] == "LINE"
        ]
        text = "\n".join(lines)
        logger.info(f"Textract extracted {len(text)} chars from {filename}")
        return text
    except Exception as e:
        logger.warning(f"Textract failed for {filename}: {e} — falling back to local extraction")
        return ""


def _extract_from_pdf_textract(file_bytes: bytes, filename: str) -> str:
    """Try Textract first; fall back to PyPDF2 if Textract fails."""
    text = _textract_bytes(file_bytes, filename)
    if text.strip():
        return text[:10000]

    # Fallback: PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        parts = [page.extract_text() or "" for page in reader.pages]
        result = "\n".join(p for p in parts if p.strip())
        logger.info(f"PyPDF2 fallback extracted {len(result)} chars from {filename}")
        return result[:10000]
    except Exception as e:
        logger.error(f"PyPDF2 fallback also failed: {e}")
        return ""


def _extract_from_image_textract(file_bytes: bytes, filename: str) -> str:
    """Try Textract first; fall back to pytesseract if Textract fails."""
    text = _textract_bytes(file_bytes, filename)
    if text.strip():
        return text[:10000]

    # Fallback: pytesseract
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(img)
        logger.info(f"pytesseract fallback extracted {len(text)} chars from {filename}")
        return text[:10000]
    except Exception as e:
        logger.warning(f"pytesseract fallback failed: {e}")
        return (
            "[Image uploaded — text extraction failed. "
            "Ensure AWS credentials are valid or Tesseract is installed.]"
        )


# ── Local extractors ──────────────────────────────────────────────────────────

def _extract_from_csv(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        reader = csv.DictReader(io.StringIO(text))
        rows = list(reader)
        headers = reader.fieldnames or []
        if not rows:
            return text[:6000]

        logger.info(f"CSV: {len(rows)} rows, {len(headers)} columns")

        numeric_cols: dict[str, float] = {}
        string_samples: dict[str, list] = {}

        for col in headers:
            values = [row.get(col, "").strip() for row in rows if row.get(col, "").strip()]
            numeric_vals = []
            for v in values:
                try:
                    cleaned = v.replace(",", "").replace("$", "").replace("£", "").replace("€", "").strip()
                    numeric_vals.append(float(cleaned))
                except ValueError:
                    pass
            if numeric_vals and len(numeric_vals) > len(values) * 0.5:
                numeric_cols[col] = sum(numeric_vals)
            else:
                unique_vals = list(dict.fromkeys(v for v in values if v))[:5]
                if unique_vals:
                    string_samples[col] = unique_vals

        summary = [
            f"CSV File Summary ({len(rows)} rows, {len(headers)} columns)",
            f"Columns: {', '.join(headers)}", "",
        ]
        if string_samples:
            summary.append("Column values (samples):")
            for col, vals in string_samples.items():
                summary.append(f"  {col}: {', '.join(str(v) for v in vals)}")
            summary.append("")
        if numeric_cols:
            summary.append("Numeric column totals:")
            for col, total in numeric_cols.items():
                summary.append(f"  {col}: {total:,.4f}")
            summary.append("")
        summary.append("Sample rows (first 5):")
        for row in rows[:5]:
            row_str = " | ".join(f"{k}: {v}" for k, v in row.items() if v.strip())
            summary.append(f"  {row_str}")
        return "\n".join(summary)
    except Exception as e:
        logger.error(f"CSV processing failed: {e}")
        return file_bytes.decode("utf-8", errors="replace")[:6000]


def _extract_from_excel(file_bytes: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames[:3]:
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [str(c) if c is not None else "" for c in rows[0]]
            parts.append(" | ".join(headers))
            for row in rows[1:51]:
                parts.append(" | ".join(str(c) if c is not None else "" for c in row))
            if len(rows) > 51:
                parts.append(f"... ({len(rows) - 51} more rows)")
        return "\n".join(parts)[:8000]
    except Exception as e:
        logger.error(f"Excel extraction failed: {e}")
        return ""


def _extract_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)[:8000]
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""
